"""Convert Markdown report to formatted Word document using python-docx."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

# ---------------------------------------------------------------------------
# Formatting constants
# ---------------------------------------------------------------------------

_FONT_NAME = "Arial"
_BODY_SIZE = Pt(11)
_HEADING1_SIZE = Pt(14)
_HEADING2_SIZE = Pt(12)
_HEADING3_SIZE = Pt(11)
_TABLE_SIZE = Pt(10)
_NOTE_SIZE = Pt(10)

# Paragraph spacing
_BODY_SPACE_AFTER = Pt(6)
_BULLET_SPACE_AFTER = Pt(3)

# Heading spacing (before / after)
_H1_SPACE_BEFORE = Pt(18)
_H1_SPACE_AFTER = Pt(12)
_H2_SPACE_BEFORE = Pt(12)
_H2_SPACE_AFTER = Pt(6)
_H3_SPACE_BEFORE = Pt(6)
_H3_SPACE_AFTER = Pt(6)

# Page layout
_MARGIN = Cm(2.54)
_BULLET_INDENT = Cm(0.63)
_BULLET_HANGING = Cm(0.32)

# Table formatting
_HEADER_SHADING = "D5E8F0"  # Light blue for table headers
_BORDER_COLOR = "BFBFBF"  # Grey for table borders
_CELL_PAD_TOP_PT = 4
_CELL_PAD_BOTTOM_PT = 4
_CELL_PAD_LEFT_PT = 6
_CELL_PAD_RIGHT_PT = 6

# ---------------------------------------------------------------------------
# Regex patterns
# ---------------------------------------------------------------------------

_HTML_COMMENT_RE = re.compile(r"<!--.*?-->", re.DOTALL)
_HEADING_RE = re.compile(r"^(#{1,4})\s+(.+)$")
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_ITALIC_RE = re.compile(r"(?<!\*)\*([^*]+?)\*(?!\*)")
_TABLE_ROW_RE = re.compile(r"^\|(.+)\|$")
_TABLE_SEP_RE = re.compile(r"^\|[\s\-:|]+\|$")
_BULLET_RE = re.compile(r"^[-*]\s+(.+)$")
_BLOCKQUOTE_RE = re.compile(r"^>\s*(.*)$")
_HLINE_RE = re.compile(r"^---+$")
_SECTION_HEADING_RE = re.compile(r"^Section\s+(\d+):\s+(.+)$")


# ---------------------------------------------------------------------------
# Markdown pre-processing for docx
# ---------------------------------------------------------------------------


def _preprocess_for_docx(markdown: str) -> str:
    """Pre-process markdown before docx conversion.

    - Converts report title from H1 to bold paragraph
    - Strips Confidence Summary section and Items Requiring Review
    - Transforms 'Section N: Title' headings to 'N. Title'
    """
    lines = markdown.split("\n")
    result: list[str] = []
    in_confidence = False

    for line in lines:
        stripped = line.strip()

        # Convert report title from heading to bold paragraph
        if stripped.startswith("# ") and "PSI Report Sections" in stripped:
            title_text = stripped.lstrip("# ").strip()
            result.append(f"**{title_text}**")
            continue

        # Start skipping at Confidence Summary heading
        if re.match(r"^##\s+Confidence\s+Summary", stripped):
            in_confidence = True
            continue

        # Stop skipping at next ## heading (that isn't part of confidence)
        if in_confidence and stripped.startswith("## "):
            in_confidence = False
            # Remove trailing --- and empty lines from before confidence section
            while result and result[-1].strip() in ("---", ""):
                result.pop()
            # Fall through to process this line

        # Skip lines within the confidence section
        if in_confidence:
            continue

        # Transform "## Section N: Title" to "## N. Title"
        heading_match = _HEADING_RE.match(stripped)
        if heading_match:
            hashes = heading_match.group(1)
            heading_text = heading_match.group(2)
            section_match = _SECTION_HEADING_RE.match(heading_text)
            if section_match:
                num = section_match.group(1)
                title = section_match.group(2)
                result.append(f"{hashes} {num}. {title}")
                continue

        result.append(line)

    return "\n".join(result)


# ---------------------------------------------------------------------------
# DocxExporter
# ---------------------------------------------------------------------------


class DocxExporter:
    """Convert Markdown report to formatted .docx using python-docx.

    This is a DRAFT output — focuses on structural correctness (headings,
    tables, content) rather than pixel-perfect layout.
    """

    def export(self, markdown_text: str, output_path: str | Path) -> Path:
        """Convert Markdown string to a Word document.

        Args:
            markdown_text: The full Markdown report.
            output_path: Where to save the .docx file.

        Returns:
            The output path.
        """
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        self._setup_document(doc)

        # Pre-process for docx (strip confidence, fix headings, etc.)
        cleaned = _preprocess_for_docx(markdown_text)

        # Strip HTML comments (confidence tags, placeholders)
        cleaned = _HTML_COMMENT_RE.sub("", cleaned)

        lines = cleaned.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # Skip empty lines
            if not stripped:
                i += 1
                continue

            # Horizontal rule — skip
            if _HLINE_RE.match(stripped):
                i += 1
                continue

            # Heading
            heading_match = _HEADING_RE.match(stripped)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2)
                self._add_heading(doc, text, level)
                i += 1
                continue

            # Blockquote (used for the note line — render in 10pt italic)
            bq_match = _BLOCKQUOTE_RE.match(stripped)
            if bq_match:
                text = bq_match.group(1)
                p = doc.add_paragraph()
                p.style = "Normal"
                p.paragraph_format.space_after = _BODY_SPACE_AFTER
                p.paragraph_format.line_spacing = 1.0
                self._add_formatted_runs(p, text, italic=True, font_size=_NOTE_SIZE)
                i += 1
                continue

            # Table (collect all contiguous table rows)
            if _TABLE_ROW_RE.match(stripped):
                table_lines = []
                while i < len(lines) and _TABLE_ROW_RE.match(lines[i].strip()):
                    table_lines.append(lines[i].strip())
                    i += 1
                self._add_table(doc, table_lines)
                continue

            # Bullet list
            bullet_match = _BULLET_RE.match(stripped)
            if bullet_match:
                text = bullet_match.group(1)
                p = doc.add_paragraph()
                p.style = "List Bullet"
                p.paragraph_format.space_after = _BULLET_SPACE_AFTER
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.left_indent = _BULLET_INDENT
                p.paragraph_format.first_line_indent = -_BULLET_HANGING
                self._add_formatted_runs(p, text)
                i += 1
                continue

            # Normal paragraph
            p = doc.add_paragraph()
            p.paragraph_format.space_after = _BODY_SPACE_AFTER
            p.paragraph_format.line_spacing = 1.0
            self._add_formatted_runs(p, stripped)
            i += 1

        doc.save(str(output_path))
        return output_path

    def _setup_document(self, doc: Document) -> None:
        """Configure page size, margins, and default font."""
        section = doc.sections[0]
        section.page_width = Mm(210)  # A4
        section.page_height = Mm(297)
        section.orientation = WD_ORIENT.PORTRAIT
        section.top_margin = _MARGIN
        section.bottom_margin = _MARGIN
        section.left_margin = _MARGIN
        section.right_margin = _MARGIN

        # Default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = _FONT_NAME
        font.size = _BODY_SIZE

    def _add_heading(self, doc: Document, text: str, level: int) -> None:
        """Add a heading with appropriate formatting and spacing."""
        # Map markdown levels to docx levels (#### = Heading 3 in docx)
        docx_level = min(level, 3)
        heading = doc.add_heading(text, level=docx_level)

        config = {
            1: (_HEADING1_SIZE, _H1_SPACE_BEFORE, _H1_SPACE_AFTER),
            2: (_HEADING2_SIZE, _H2_SPACE_BEFORE, _H2_SPACE_AFTER),
            3: (_HEADING3_SIZE, _H3_SPACE_BEFORE, _H3_SPACE_AFTER),
        }
        size, space_before, space_after = config.get(docx_level, config[3])

        heading.paragraph_format.space_before = space_before
        heading.paragraph_format.space_after = space_after

        for run in heading.runs:
            run.font.name = _FONT_NAME
            run.font.size = size
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)

    def _add_formatted_runs(
        self,
        paragraph,
        text: str,
        italic: bool = False,
        font_size=None,
    ) -> None:
        """Parse inline formatting (**bold**, *italic*) and add runs."""
        size = font_size or _BODY_SIZE
        parts = _split_formatted(text)
        for content, is_bold, is_italic in parts:
            run = paragraph.add_run(content)
            run.font.name = _FONT_NAME
            run.font.size = size
            run.bold = is_bold
            run.italic = is_italic or italic

    def _add_table(self, doc: Document, table_lines: list[str]) -> None:
        """Convert pipe-delimited lines to a Word table with professional formatting."""
        # Parse rows, skip separator lines
        rows: list[list[str]] = []
        for line in table_lines:
            if _TABLE_SEP_RE.match(line):
                continue
            cells = [c.strip() for c in line.strip("|").split("|")]
            rows.append(cells)

        if not rows:
            return

        num_cols = max(len(r) for r in rows)
        table = doc.add_table(rows=len(rows), cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set table to full page width
        _set_table_full_width(table)

        # Set even column widths across available space
        available_cm = 21.0 - 2 * 2.54  # A4 width minus margins
        col_width = Cm(available_cm / num_cols)
        for col in table.columns:
            col.width = col_width

        for ri, row_data in enumerate(rows):
            for ci in range(num_cols):
                cell = table.cell(ri, ci)
                text = row_data[ci] if ci < len(row_data) else ""

                # Clear default paragraph and add formatted text
                p = cell.paragraphs[0]
                p.clear()
                self._add_formatted_runs(p, text, font_size=_TABLE_SIZE)

                # Header row: bold + shading
                if ri == 0:
                    for run in p.runs:
                        run.bold = True
                    _set_cell_shading(cell, _HEADER_SHADING)

                # Cell borders and padding
                _set_cell_border(cell)
                _set_cell_margins(
                    cell,
                    _CELL_PAD_TOP_PT,
                    _CELL_PAD_BOTTOM_PT,
                    _CELL_PAD_LEFT_PT,
                    _CELL_PAD_RIGHT_PT,
                )

        # Space after table
        doc.add_paragraph()


# ---------------------------------------------------------------------------
# Inline formatting parser
# ---------------------------------------------------------------------------


def _split_formatted(text: str) -> list[tuple[str, bool, bool]]:
    """Split text into (content, is_bold, is_italic) tuples.

    Handles **bold** and *italic* markers.
    """
    result: list[tuple[str, bool, bool]] = []
    pos = 0

    while pos < len(text):
        # Check for bold
        bold_match = _BOLD_RE.search(text, pos)
        italic_match = _ITALIC_RE.search(text, pos)

        next_match = None
        is_bold = False
        is_italic = False

        if bold_match and (not italic_match or bold_match.start() <= italic_match.start()):
            next_match = bold_match
            is_bold = True
        elif italic_match:
            next_match = italic_match
            is_italic = True

        if next_match is None:
            # No more formatting — add rest as plain text
            if pos < len(text):
                result.append((text[pos:], False, False))
            break

        # Add plain text before the match
        if next_match.start() > pos:
            result.append((text[pos:next_match.start()], False, False))

        # Add the formatted text
        result.append((next_match.group(1), is_bold, is_italic))
        pos = next_match.end()

    return result if result else [("", False, False)]


# ---------------------------------------------------------------------------
# Table XML helpers
# ---------------------------------------------------------------------------


def _set_table_full_width(table) -> None:
    """Set table preferred width to 100% of page width."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = tbl.makeelement(qn("w:tblPr"), {})
        tbl.insert(0, tblPr)
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = tbl.makeelement(qn("w:tblW"), {})
        tblPr.append(tblW)
    # 5000 = 100% in OOXML pct units (fiftieths of a percent)
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")


def _set_cell_border(cell) -> None:
    """Set thin grey borders on a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove existing borders if any
    existing = tcPr.find(qn("w:tcBorders"))
    if existing is not None:
        tcPr.remove(existing)
    tcBorders = tc.makeelement(qn("w:tcBorders"), {})
    tcPr.append(tcBorders)

    for edge in ("top", "left", "bottom", "right"):
        element = tc.makeelement(
            qn(f"w:{edge}"),
            {
                qn("w:val"): "single",
                qn("w:sz"): "4",
                qn("w:space"): "0",
                qn("w:color"): _BORDER_COLOR,
            },
        )
        tcBorders.append(element)


def _set_cell_shading(cell, color: str) -> None:
    """Set background shading on a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = tc.makeelement(
        qn("w:shd"),
        {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): color,
        },
    )
    tcPr.append(shading)


def _set_cell_margins(cell, top_pt: int, bottom_pt: int, left_pt: int, right_pt: int) -> None:
    """Set cell padding/margins in points."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove existing margins if any
    existing = tcPr.find(qn("w:tcMar"))
    if existing is not None:
        tcPr.remove(existing)
    tcMar = tc.makeelement(qn("w:tcMar"), {})
    for edge, value_pt in [("top", top_pt), ("bottom", bottom_pt), ("left", left_pt), ("right", right_pt)]:
        # Convert points to twips (20 twips per point)
        twips = int(value_pt * 20)
        el = tc.makeelement(
            qn(f"w:{edge}"),
            {
                qn("w:w"): str(twips),
                qn("w:type"): "dxa",
            },
        )
        tcMar.append(el)
    tcPr.append(tcMar)
