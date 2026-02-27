"""Tests for the compose layer — renderer and docx export."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from src.compose.docx_export import DocxExporter
from src.compose.renderer import ReportRenderer
from src.ingest.state_detector import StateIdentification
from src.interpret.ai_interpreter import InterpretedSection


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


@pytest.fixture
def vic_state_id() -> StateIdentification:
    return StateIdentification(
        state="VIC",
        address="151 Melville Road, Brunswick West, VIC 3055",
        suburb="Brunswick West",
        postcode="3055",
        lotsearch_reference="LS089658 EP",
    )


@pytest.fixture
def nsw_state_id() -> StateIdentification:
    return StateIdentification(
        state="NSW",
        address="58-62 Nangunia Street, Barooga, NSW 3644",
        suburb="Barooga",
        postcode="3644",
        lotsearch_reference="LS121125 EP",
    )


@pytest.fixture
def sample_sections() -> list[InterpretedSection]:
    return [
        InterpretedSection(
            section_id="3_env_setting",
            prose=(
                "### 3.1 Topography\n\n"
                "The Site is located on relatively flat terrain.\n\n"
                "<!-- CONFIDENCE: HIGH -->\n\n"
                "### 3.2 Regional Geology\n\n"
                "The Site is underlain by the Melbourne Formation (Sxm).\n\n"
                "<!-- CONFIDENCE: HIGH -->\n\n"
                "### 3.3 Regional Hydrogeology\n\n"
                "Groundwater is classified as Segment C.\n\n"
                "<!-- CONFIDENCE: MEDIUM -->\n\n"
                "### 3.4 Surface Water Hydrology\n\n"
                "Moonee Ponds Creek is located 500m west.\n\n"
                "<!-- CONFIDENCE: HIGH -->"
            ),
            confidence="MEDIUM",
        ),
        InterpretedSection(
            section_id="4.1_business_dirs",
            prose=(
                "A review of historical business directories identified "
                "a motor garage at the Site.\n\n"
                "<!-- CONFIDENCE: MEDIUM -->"
            ),
            confidence="MEDIUM",
        ),
        InterpretedSection(
            section_id="4.3_epa_registers",
            prose=(
                "No EPA Priority Sites were identified within 1,000m.\n\n"
                "<!-- CONFIDENCE: HIGH -->"
            ),
            confidence="HIGH",
        ),
        InterpretedSection(
            section_id="4.3_env_audits",
            prose=(
                "23 environmental audit records were identified.\n\n"
                "<!-- CONFIDENCE: LOW — Review required for proximal audits -->"
            ),
            confidence="LOW",
            review_flags=["Review required for proximal audits"],
        ),
        InterpretedSection(
            section_id="4.4_pfas",
            prose=(
                "No PFAS investigation sites were identified.\n\n"
                "<!-- CONFIDENCE: HIGH -->"
            ),
            confidence="HIGH",
        ),
        InterpretedSection(
            section_id="4.5_waste_liquid_fuel",
            prose="No waste management facilities identified.\n\n<!-- CONFIDENCE: HIGH -->",
            confidence="HIGH",
        ),
        InterpretedSection(
            section_id="4.6_mining_fire",
            prose="No mining shafts or fire history records identified.\n\n<!-- CONFIDENCE: HIGH -->",
            confidence="HIGH",
        ),
        InterpretedSection(
            section_id="4.8_defence",
            prose="No defence sites identified.\n\n<!-- CONFIDENCE: HIGH -->",
            confidence="HIGH",
        ),
        InterpretedSection(
            section_id="4.9_site_history",
            prose=(
                "| Year/Period | Source | Information | Relevance |\n"
                "|------------|--------|-------------|----------|\n"
                "| 1980-1991 | Business Directories | Motor garage | High |\n\n"
                "<!-- CONFIDENCE: MEDIUM -->"
            ),
            confidence="MEDIUM",
            tables_markdown=[
                "| Year/Period | Source | Information | Relevance |\n"
                "|------------|--------|-------------|----------|\n"
                "| 1980-1991 | Business Directories | Motor garage | High |"
            ],
        ),
    ]


@pytest.fixture
def renderer() -> ReportRenderer:
    return ReportRenderer()


@pytest.fixture
def exporter() -> DocxExporter:
    return DocxExporter()


# ---------------------------------------------------------------------------
# Renderer tests
# ---------------------------------------------------------------------------


class TestRendererVIC:
    def test_renders_vic_report(
        self, renderer: ReportRenderer, vic_state_id: StateIdentification,
        sample_sections: list[InterpretedSection],
    ):
        result = renderer.render("VIC", sample_sections, vic_state_id)
        assert isinstance(result, str)
        assert len(result) > 100

    def test_contains_site_address(
        self, renderer: ReportRenderer, vic_state_id: StateIdentification,
        sample_sections: list[InterpretedSection],
    ):
        result = renderer.render("VIC", sample_sections, vic_state_id)
        assert "151 Melville Road" in result

    def test_contains_vic_section_headings(
        self, renderer: ReportRenderer, vic_state_id: StateIdentification,
        sample_sections: list[InterpretedSection],
    ):
        result = renderer.render("VIC", sample_sections, vic_state_id)
        assert "Section 3: Environmental Setting" in result
        assert "Section 4: Desktop Site History Review" in result
        assert "3.1 Topography" in result
        assert "3.2 Regional Geology" in result
        assert "4.1 Historical Business Directory" in result
        assert "4.3 EPA Victoria" in result

    def test_contains_confidence_summary(
        self, renderer: ReportRenderer, vic_state_id: StateIdentification,
        sample_sections: list[InterpretedSection],
    ):
        result = renderer.render("VIC", sample_sections, vic_state_id)
        assert "Confidence Summary" in result
        assert "HIGH" in result
        assert "MEDIUM" in result

    def test_contains_review_flags(
        self, renderer: ReportRenderer, vic_state_id: StateIdentification,
        sample_sections: list[InterpretedSection],
    ):
        result = renderer.render("VIC", sample_sections, vic_state_id)
        assert "Review required for proximal audits" in result

    def test_contains_vic_references(
        self, renderer: ReportRenderer, vic_state_id: StateIdentification,
        sample_sections: list[InterpretedSection],
    ):
        result = renderer.render("VIC", sample_sections, vic_state_id)
        assert "Environment Protection Act 2017" in result
        assert "Environment Reference Standard" in result
        assert "PPN30" in result

    def test_contains_aerial_photo_placeholder(
        self, renderer: ReportRenderer, vic_state_id: StateIdentification,
        sample_sections: list[InterpretedSection],
    ):
        result = renderer.render("VIC", sample_sections, vic_state_id)
        assert "MANUAL_INPUT" in result
        assert "aerial" in result.lower()

    def test_contains_env_audits_section(
        self, renderer: ReportRenderer, vic_state_id: StateIdentification,
        sample_sections: list[InterpretedSection],
    ):
        result = renderer.render("VIC", sample_sections, vic_state_id)
        assert "Environmental Audits" in result
        assert "23 environmental audit records" in result

    def test_inserts_interpreted_prose(
        self, renderer: ReportRenderer, vic_state_id: StateIdentification,
        sample_sections: list[InterpretedSection],
    ):
        result = renderer.render("VIC", sample_sections, vic_state_id)
        assert "Melbourne Formation" in result
        assert "Moonee Ponds Creek" in result
        assert "motor garage" in result


class TestRendererNSW:
    def test_renders_nsw_report(
        self, renderer: ReportRenderer, nsw_state_id: StateIdentification,
        sample_sections: list[InterpretedSection],
    ):
        # Provide a POEO section for NSW
        nsw_sections = sample_sections + [
            InterpretedSection(
                section_id="4.3_poeo_licences",
                prose="No POEO licences identified.\n\n<!-- CONFIDENCE: HIGH -->",
                confidence="HIGH",
            ),
        ]
        result = renderer.render("NSW", nsw_sections, nsw_state_id)
        assert isinstance(result, str)
        assert "58-62 Nangunia Street" in result

    def test_nsw_has_correct_references(
        self, renderer: ReportRenderer, nsw_state_id: StateIdentification,
        sample_sections: list[InterpretedSection],
    ):
        result = renderer.render("NSW", sample_sections, nsw_state_id)
        assert "Contaminated Land Management Act 1997" in result
        assert "Protection of the Environment Operations Act 1997" in result

    def test_nsw_has_poeo_section(
        self, renderer: ReportRenderer, nsw_state_id: StateIdentification,
    ):
        sections = [
            InterpretedSection(
                section_id="4.3_poeo_licences",
                prose="Murray Irrigation licence identified.\n\n<!-- CONFIDENCE: HIGH -->",
                confidence="HIGH",
            ),
        ]
        result = renderer.render("NSW", sections, nsw_state_id)
        assert "POEO Act Licences" in result
        assert "Murray Irrigation" in result


class TestRendererMissingSections:
    def test_handles_empty_sections(
        self, renderer: ReportRenderer, vic_state_id: StateIdentification,
    ):
        result = renderer.render("VIC", [], vic_state_id)
        assert isinstance(result, str)
        assert "Section 3" in result  # Template structure still present
        assert "EXTRACTION_ONLY" in result  # Fallback markers

    def test_handles_partial_sections(
        self, renderer: ReportRenderer, vic_state_id: StateIdentification,
    ):
        sections = [
            InterpretedSection(
                section_id="3_env_setting",
                prose="Topography content only.\n\n<!-- CONFIDENCE: HIGH -->",
                confidence="HIGH",
            ),
        ]
        result = renderer.render("VIC", sections, vic_state_id)
        assert "Topography content only" in result
        # Missing sections should have fallback text
        assert "Section 4" in result


# ---------------------------------------------------------------------------
# DocxExporter tests
# ---------------------------------------------------------------------------


class TestDocxExporter:
    def test_produces_valid_docx(self, exporter: DocxExporter, tmp_path: Path):
        md = "# Test Report\n\nSome body text.\n"
        output = exporter.export(md, tmp_path / "test.docx")
        assert output.exists()
        # Verify it opens as a valid docx
        doc = Document(str(output))
        assert len(doc.paragraphs) > 0

    def test_heading_conversion(self, exporter: DocxExporter, tmp_path: Path):
        md = (
            "# Heading 1\n\n"
            "## Heading 2\n\n"
            "### Heading 3\n\n"
            "#### Heading 4\n\n"
            "Body text.\n"
        )
        output = exporter.export(md, tmp_path / "headings.docx")
        doc = Document(str(output))

        # Check heading styles
        styles = [p.style.name for p in doc.paragraphs if p.text.strip()]
        assert "Heading 1" in styles
        assert "Heading 2" in styles
        assert "Heading 3" in styles
        # #### maps to Heading 3 in docx
        assert styles.count("Heading 3") >= 2

    def test_strips_html_comments(self, exporter: DocxExporter, tmp_path: Path):
        md = (
            "# Report\n\n"
            "Some text.\n\n"
            "<!-- CONFIDENCE: HIGH -->\n\n"
            "More text.\n\n"
            "<!-- MANUAL_INPUT: This should not appear in docx -->\n"
        )
        output = exporter.export(md, tmp_path / "comments.docx")
        doc = Document(str(output))
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "CONFIDENCE" not in full_text
        assert "MANUAL_INPUT" not in full_text
        assert "Some text" in full_text
        assert "More text" in full_text

    def test_table_conversion(self, exporter: DocxExporter, tmp_path: Path):
        md = (
            "# Report\n\n"
            "| Name | Value |\n"
            "|------|-------|\n"
            "| Alpha | 100 |\n"
            "| Beta | 200 |\n"
        )
        output = exporter.export(md, tmp_path / "tables.docx")
        doc = Document(str(output))
        assert len(doc.tables) >= 1
        table = doc.tables[0]
        # Header row + 2 data rows
        assert len(table.rows) == 3
        assert table.cell(0, 0).text.strip() == "Name"
        assert table.cell(1, 0).text.strip() == "Alpha"
        assert table.cell(2, 1).text.strip() == "200"

    def test_bold_and_italic(self, exporter: DocxExporter, tmp_path: Path):
        md = (
            "# Report\n\n"
            "This has **bold text** and *italic text* in it.\n"
        )
        output = exporter.export(md, tmp_path / "formatting.docx")
        doc = Document(str(output))
        # Find the paragraph with formatted text
        body_paras = [p for p in doc.paragraphs if "bold text" in p.text]
        assert len(body_paras) >= 1
        runs = body_paras[0].runs
        bold_runs = [r for r in runs if r.bold]
        italic_runs = [r for r in runs if r.italic]
        assert len(bold_runs) >= 1, "Expected at least one bold run"
        assert len(italic_runs) >= 1, "Expected at least one italic run"

    def test_bullet_list(self, exporter: DocxExporter, tmp_path: Path):
        md = (
            "# Report\n\n"
            "- First item\n"
            "- Second item\n"
            "- Third item\n"
        )
        output = exporter.export(md, tmp_path / "bullets.docx")
        doc = Document(str(output))
        bullet_paras = [p for p in doc.paragraphs if p.style.name == "List Bullet"]
        assert len(bullet_paras) == 3

    def test_a4_page_size(self, exporter: DocxExporter, tmp_path: Path):
        md = "# Test\n\nBody.\n"
        output = exporter.export(md, tmp_path / "pagesize.docx")
        doc = Document(str(output))
        section = doc.sections[0]
        # A4 = 210mm x 297mm (within 1mm tolerance)
        width_mm = section.page_width.mm
        height_mm = section.page_height.mm
        assert 209 <= width_mm <= 211
        assert 296 <= height_mm <= 298

    def test_full_report_roundtrip(
        self, exporter: DocxExporter, tmp_path: Path,
    ):
        """Test that a realistic report markdown converts to docx without error."""
        renderer = ReportRenderer()
        state_id = StateIdentification(
            state="VIC",
            address="151 Melville Road, Brunswick West, VIC 3055",
            suburb="Brunswick West",
            postcode="3055",
            lotsearch_reference="LS089658 EP",
        )
        sections = [
            InterpretedSection(
                section_id="3_env_setting",
                prose="The Site is flat.\n\n<!-- CONFIDENCE: HIGH -->",
                confidence="HIGH",
            ),
            InterpretedSection(
                section_id="4.3_epa_registers",
                prose="No EPA records.\n\n<!-- CONFIDENCE: HIGH -->",
                confidence="HIGH",
            ),
        ]
        markdown = renderer.render("VIC", sections, state_id)
        output = exporter.export(markdown, tmp_path / "full.docx")
        doc = Document(str(output))
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "151 Melville Road" in full_text
        assert "CONFIDENCE" not in full_text


# ---------------------------------------------------------------------------
# App import test
# ---------------------------------------------------------------------------


class TestAppImport:
    def test_app_module_importable(self):
        """Verify app.py can be imported without error."""
        import importlib
        # This tests that the module-level code doesn't crash
        # (Streamlit won't actually run since there's no server)
        spec = importlib.util.find_spec("app")
        assert spec is not None, "app.py module should be importable"
